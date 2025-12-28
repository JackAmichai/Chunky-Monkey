"""
DOCX (Word) document parser.

Extracts text and structure from Microsoft Word documents.
Requires python-docx library.
"""

from __future__ import annotations

from pathlib import Path
from typing import BinaryIO

from monkey.parsers.base import BaseParser, ParsedDocument, DocumentElement


class DocxParser(BaseParser):
    """
    Parse Microsoft Word (.docx) documents into structured elements.
    
    Extracts headings, paragraphs, tables, and lists while preserving
    document structure.
    
    Requires: pip install python-docx
    Or: pip install chunky-monkey[docx]
    
    Example:
        >>> parser = DocxParser()
        >>> doc = parser.parse_file("document.docx")
        >>> for element in doc.elements:
        ...     print(element.type, element.content[:50])
    """
    
    # Word heading styles
    HEADING_STYLES = {
        "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
        "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
        "Title": 1, "Subtitle": 2,
    }
    
    def __init__(
        self,
        extract_tables: bool = True,
        extract_images: bool = False,
        include_headers_footers: bool = False,
    ):
        """
        Initialize DOCX parser.
        
        Args:
            extract_tables: Extract tables as structured elements
            extract_images: Include image placeholders in output
            include_headers_footers: Include header/footer content
        """
        self.extract_tables = extract_tables
        self.extract_images = extract_images
        self.include_headers_footers = include_headers_footers
        self._docx = None
    
    def _get_docx(self):
        """Import and return docx module."""
        if self._docx is None:
            try:
                import docx
                self._docx = docx
            except ImportError:
                raise ImportError(
                    "python-docx is required for DOCX parsing.\n"
                    "Install with: pip install python-docx\n"
                    "Or: pip install chunky-monkey[docx]"
                )
        return self._docx
    
    def parse(self, text: str) -> ParsedDocument:
        """
        Parse from text (not typical for DOCX - use parse_file).
        
        For DOCX files, use parse_file() or parse_bytes() instead.
        This method treats the input as pre-extracted text.
        """
        elements = []
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        
        position = 0
        for para in paragraphs:
            elements.append(DocumentElement(
                type="paragraph",
                content=para,
                position=(position, position + len(para)),
            ))
            position += len(para) + 2
        
        return ParsedDocument(elements=elements)
    
    def parse_file(self, path: str | Path) -> ParsedDocument:
        """
        Parse a DOCX file.
        
        Args:
            path: Path to DOCX file
            
        Returns:
            ParsedDocument with extracted elements
        """
        docx = self._get_docx()
        path = Path(path)
        
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")
        
        doc = docx.Document(str(path))
        return self._parse_document(doc, source=str(path))
    
    def parse_bytes(self, data: bytes, source: str | None = None) -> ParsedDocument:
        """
        Parse DOCX from bytes.
        
        Args:
            data: DOCX file content as bytes
            source: Optional source identifier
            
        Returns:
            ParsedDocument with extracted elements
        """
        import io
        docx = self._get_docx()
        doc = docx.Document(io.BytesIO(data))
        return self._parse_document(doc, source=source)
    
    def _parse_document(self, doc, source: str | None = None) -> ParsedDocument:
        """Parse a python-docx Document object."""
        elements = []
        position = 0
        title = None
        current_section = None
        
        # Extract core properties for title
        if doc.core_properties.title:
            title = doc.core_properties.title
        
        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Determine element type from style
            style_name = para.style.name if para.style else ""
            
            if style_name in self.HEADING_STYLES:
                level = self.HEADING_STYLES[style_name]
                current_section = text
                
                # First heading might be title
                if title is None and level == 1:
                    title = text
                
                elements.append(DocumentElement(
                    type="header",
                    content=text,
                    level=level,
                    position=(position, position + len(text)),
                    section_header=current_section,
                ))
            elif style_name.startswith("List"):
                # List items
                elements.append(DocumentElement(
                    type="list_item",
                    content=text,
                    position=(position, position + len(text)),
                    section_header=current_section,
                ))
            else:
                # Regular paragraph
                elements.append(DocumentElement(
                    type="paragraph",
                    content=text,
                    position=(position, position + len(text)),
                    section_header=current_section,
                ))
            
            position += len(text) + 2
        
        # Extract tables
        if self.extract_tables:
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    elements.append(DocumentElement(
                        type="table",
                        content=table_text,
                        position=(position, position + len(table_text)),
                        section_header=current_section,
                    ))
                    position += len(table_text) + 2
        
        return ParsedDocument(
            elements=elements,
            source=source,
            title=title,
            metadata={
                "author": doc.core_properties.author,
                "created": str(doc.core_properties.created) if doc.core_properties.created else None,
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None,
            }
        )
    
    def _extract_table(self, table) -> str:
        """Extract table as markdown-style text."""
        rows = []
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip().replace("\n", " "))
            rows.append(cells)
        
        if not rows:
            return ""
        
        # Build markdown table
        lines = []
        for i, row in enumerate(rows):
            lines.append("| " + " | ".join(row) + " |")
            if i == 0:
                lines.append("|" + "|".join(["---"] * len(row)) + "|")
        
        return "\n".join(lines)
